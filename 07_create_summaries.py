import pandas as pd
import numpy as np
import tiktoken
from model_utils import *
from random import shuffle
from docx import Document
import os

cur_directory = os.path.dirname(os.path.abspath(__file__))
relative_input_directory = 'outputs/intermediate/'
relative_finaloutput_directory = 'outputs/final/'

input_path = os.path.join(cur_directory, relative_input_directory)
final_output_path = os.path.join(cur_directory, relative_finaloutput_directory)

df = pd.read_excel(input_path + "topic_results.xlsx")

topic_set = ["Standards of Conduct",
            "Duration for Early Head Start",
            "Duration for Head Start Preschool",
            "Lead in Water",
            "Lead in Paint",
            "Governing body and parent committees",
            "Health, nutrition and mental health",
            "Data Privacy and FERPA",
            "Program Structure: class size, hours of operation, length of school year",
            "Center or home based program structure",
            "Disabilities, IDEA, 504 Plan and Individualized Education Plan",
            "Delegate Agencies",
            "Funding suspension and refunding appeals",
            "<...>" # add additional topics here
            ]

encoding = tiktoken.get_encoding("cl100k_base")

def att_token_length(row):
    if pd.isna(row) or isinstance(row, int):
        return 0
    return len(encoding.encode(row))

df['tokens'] = df.apply(lambda x: att_token_length(x['segment']), axis=1)

topic_dict = {}

for topic in topic_set:
    df_topic = df[df['topic_clean']== topic].sort_values(['text_chunk_id','segment_number_min'])
    doc_ids = df_topic.Document.unique()
    doc_ids_sample = []
    if df_topic.tokens.sum() > 70000:
        shuffle(doc_ids)
        doc_ids_sample.append(doc_ids[0])
        for i in doc_ids[1:]:
            doc_tokens = df_topic[df_topic['Document']==i]['tokens'].sum()
            df_topic_sample = df_topic[df_topic['Document'].isin(doc_ids_sample)]
            all_tokens = df_topic_sample.tokens.sum()
            if all_tokens + doc_tokens > 70000:
                break
            doc_ids_sample.append(i)
        df_topic_concat = df_topic_sample.groupby(["Document","topic_clean"])['segment'].apply("\n".join).reset_index(name='segment_concat') 
    else:
        df_topic_concat = df_topic.groupby(["Document","topic_clean"])['segment'].apply("\n".join).reset_index(name='segment_concat')       
    topic_dict[topic] = topic_dict.get(topic, df_topic_concat)
    df_topic_concat.to_excel(final_output_path + "{}_summary_segments.xlsx".format(topic))


lang4 = modelType('prod_gpt4')

summarize_dict = {}
for topic in topic_set:
    prompt = '''You are reading a list of segments of comments written by the public related to a rule 
            change coming from the office of Head Start. Summarize what commenters are saying about the topic
            delineated by triple back ticks in about 500 words. The list of segments is contained in square brackets.'''
    full_prompt = prompt + " ```{}``` ".format(topic) + str(list(topic_dict[topic]['segment_concat']))
    response = call_azure_openai('1', full_prompt, lang4)

    with open('{}.txt'.format(topic), 'w') as f:
        f.write(response)

    summarize_dict[topic] = summarize_dict.get(topic, response)

summary_df = pd.DataFrame([summarize_dict]).transpose()
summary_df.reset_index(inplace=True)
summary_df.columns = ['topic','summary']

summary_df.to_excel(final_output_path + "summary_df.xlsx")


def format_dataframe_to_word(df):
    # Create a new Word document
    for top in df.topic.unique():
        doc = Document()
        # Add Heading 1 for the topic
        summary =  df.loc[df['topic']==top, 'summary'].values
        doc.add_heading(top, level=1)
        doc.add_paragraph(summary)

        # Save the document
        doc.save(final_output_path + "{}_summary.docx".format(top))

# Replace 'output_document.docx' with the desired output Word document file name

format_dataframe_to_word(summary_df)

print("summaries completed")
